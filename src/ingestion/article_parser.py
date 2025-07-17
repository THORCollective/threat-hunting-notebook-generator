"""Article parsing and text extraction module."""

import os
from pathlib import Path
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import requests


class ArticleParser:
    """Extract text content from various article formats."""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.md']
    
    def parse_article(self, source: str) -> Dict[str, Any]:
        """
        Parse article from file path or URL.
        
        Args:
            source: File path or URL to the article
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        if source.startswith(('http://', 'https://')):
            return self._parse_url(source)
        else:
            return self._parse_file(source)
    
    def _parse_file(self, file_path: str) -> Dict[str, Any]:
        """Parse article from local file."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return self._parse_pdf(file_path)
        elif extension == '.docx':
            return self._parse_docx(file_path)
        elif extension in ['.txt', '.md']:
            return self._parse_text(file_path)
        elif extension == '.html':
            return self._parse_html_file(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    def _parse_url(self, url: str) -> Dict[str, Any]:
        """Parse article from URL."""
        try:
            response = requests.get(url, headers={'User-Agent': 'Mozilla/5.0'})
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Extract title
            title = soup.find('title')
            title_text = title.get_text().strip() if title else "Unknown Title"
            
            # Remove script and style elements
            for element in soup(['script', 'style']):
                element.decompose()
            
            # Extract main content
            content = soup.get_text()
            
            return {
                'title': title_text,
                'content': content,
                'source': url,
                'format': 'html',
                'metadata': {
                    'url': url,
                    'content_type': response.headers.get('content-type', ''),
                }
            }
        except Exception as e:
            raise ValueError(f"Failed to parse URL {url}: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF file."""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            
            for page in reader.pages:
                text += page.extract_text() + "\n"
        
        return {
            'title': Path(file_path).stem,
            'content': text,
            'source': file_path,
            'format': 'pdf',
            'metadata': {
                'num_pages': len(reader.pages),
                'file_size': os.path.getsize(file_path)
            }
        }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return {
            'title': Path(file_path).stem,
            'content': text,
            'source': file_path,
            'format': 'docx',
            'metadata': {
                'num_paragraphs': len(doc.paragraphs),
                'file_size': os.path.getsize(file_path)
            }
        }
    
    def _parse_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
        
        return {
            'title': Path(file_path).stem,
            'content': content,
            'source': file_path,
            'format': Path(file_path).suffix[1:],
            'metadata': {
                'file_size': os.path.getsize(file_path),
                'encoding': 'utf-8'
            }
        }
    
    def _parse_html_file(self, file_path: str) -> Dict[str, Any]:
        """Extract text from HTML file."""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
        
        # Extract title
        title = soup.find('title')
        title_text = title.get_text().strip() if title else Path(file_path).stem
        
        # Remove script and style elements
        for element in soup(['script', 'style']):
            element.decompose()
        
        content = soup.get_text()
        
        return {
            'title': title_text,
            'content': content,
            'source': file_path,
            'format': 'html',
            'metadata': {
                'file_size': os.path.getsize(file_path),
                'encoding': 'utf-8'
            }
        }